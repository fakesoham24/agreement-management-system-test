import PyPDF2
import docx
import io
import gc
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def _image_to_base64(img, max_size: int = 4 * 1024 * 1024) -> str | None:
    """Convert PIL Image to base64 JPEG, progressively reducing quality/size to fit limit.
    Also ensures the image stays within the Groq Vision maximum pixel limit (33,177,600 pixels).
    Memory-optimized: reuses buffers and cleans up intermediates.
    """
    from io import BytesIO
    import base64
    import math
    from PIL import Image

    # Max pixel count constraint by Groq Vision
    max_pixels = 30000000
    if img.width * img.height > max_pixels:
        scale = math.sqrt(max_pixels / (img.width * img.height))
        new_w = int(img.width * scale)
        new_h = int(img.height * scale)
        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        logger.info(f"Image resized to {new_w}x{new_h} to fit Groq pixel count limit.")

    # Reuse a single BytesIO buffer to reduce memory allocations
    buf = BytesIO()
    for quality in (60, 40, 25):
        buf.seek(0)
        buf.truncate()
        img.save(buf, format="JPEG", quality=quality)
        raw_bytes = buf.getvalue()
        b64 = base64.b64encode(raw_bytes).decode("utf-8")
        del raw_bytes
        if len(b64) <= max_size:
            del buf
            return b64
        del b64

    # If still too large, downscale the image
    for scale in (0.5, 0.35):
        new_w = int(img.width * scale)
        new_h = int(img.height * scale)
        resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        buf.seek(0)
        buf.truncate()
        resized.save(buf, format="JPEG", quality=50)
        del resized
        raw_bytes = buf.getvalue()
        b64 = base64.b64encode(raw_bytes).decode("utf-8")
        del raw_bytes
        if len(b64) <= max_size:
            del buf
            return b64
        del b64

    del buf
    return None  # Could not fit within limit


def _ocr_page_with_vision(client, model: str, img_b64: str, page_num: int, total_pages: int) -> str:
    """Send a single page image to Groq vision model for OCR text extraction.
    
    Uses strict anti-hallucination prompting and retries on transient failures.
    """
    import time

    system_prompt = (
        "You are a precise OCR (Optical Character Recognition) engine. "
        "Your ONLY job is to read and transcribe ALL text visible in the provided image. "
        "Rules:\n"
        "1. Transcribe EVERY word, number, date, and symbol exactly as it appears.\n"
        "2. Preserve the original layout, line breaks, and paragraph structure.\n"
        "3. Do NOT add, infer, summarize, or interpret anything.\n"
        "4. Do NOT add headers, comments, or explanations.\n"
        "5. If a word is partially illegible, transcribe what you can see and mark unclear parts with [?].\n"
        "6. If the page is blank or contains no text, respond with exactly: [BLANK PAGE]\n"
        "7. Output ONLY the raw transcribed text, nothing else."
    )

    user_prompt = f"Transcribe ALL text from this scanned document image (page {page_num} of {total_pages}). Output only the raw text."

    max_retries = 2
    for attempt in range(max_retries + 1):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": user_prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}
                            }
                        ]
                    }
                ],
                temperature=0,
                max_tokens=4096
            )
            result = response.choices[0].message.content
            if result and "[BLANK PAGE]" not in result:
                return result
            return ""
        except Exception as e:
            if attempt < max_retries:
                wait = 2 ** (attempt + 1)
                logger.warning(f"Vision OCR retry {attempt+1} for page {page_num}: {e}, waiting {wait}s")
                time.sleep(wait)
            else:
                raise


def extract_text_from_scanned_pdf(file_bytes: bytes) -> str:
    """Extract text from a scanned (image-based) PDF using Groq Vision model.

    Converts each page to an image using PyMuPDF, then runs Groq Llama 4 Scout
    for text recognition.
    Falls back to PyPDF2 text extraction if OCR yields no results.

    Memory-optimized: uses 150 DPI, explicit cleanup after each page,
    and caps at 20 pages to prevent OOM on Railway.
    """
    import fitz  # PyMuPDF
    from PIL import Image
    from groq import Groq
    from backend.config import GROQ_API_KEY, GROQ_VISION_MODEL

    MAX_BASE64_SIZE = 4 * 1024 * 1024  # 4 MB Groq limit
    MAX_OCR_PAGES = 20  # Cap pages to prevent runaway memory usage

    client = Groq(api_key=GROQ_API_KEY)
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    total_pages = len(doc)
    pages_to_process = min(total_pages, MAX_OCR_PAGES)

    if total_pages > MAX_OCR_PAGES:
        logger.warning(f"PDF has {total_pages} pages, capping OCR to first {MAX_OCR_PAGES} pages.")

    try:
        for page_num in range(pages_to_process):
            pix = None
            img = None
            img_b64 = None
            try:
                page = doc.load_page(page_num)
                # Render page at 150 DPI — sufficient for Groq Vision OCR accuracy
                # while using ~4x less memory than 300 DPI
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                # Free the pixmap immediately — PIL Image has a copy of the data
                del pix
                pix = None

                # Convert PIL Image to base64 JPEG
                img_b64 = _image_to_base64(img, max_size=MAX_BASE64_SIZE)

                # Free the PIL Image — base64 string is all we need now
                del img
                img = None

                if not img_b64:
                    logger.warning(f"Page {page_num + 1}: Could not compress/downscale to under 4MB, skipping.")
                    continue

                # Run OCR using Groq Vision completion
                page_text = _ocr_page_with_vision(client, GROQ_VISION_MODEL, img_b64, page_num + 1, total_pages)

                # Free the base64 string immediately after API call
                del img_b64
                img_b64 = None

                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())

                logger.info(f"OCR processed page {page_num + 1}/{total_pages}")

            except Exception as page_err:
                logger.warning(f"OCR failed on page {page_num + 1}/{total_pages}: {page_err}")
                # Continue processing remaining pages
                continue
            finally:
                # Ensure all page-level objects are freed
                del pix, img, img_b64
                gc.collect()
    finally:
        doc.close()

    # If OCR extracted nothing, try PyPDF2 fallback
    ocr_text = "\n\n".join(text_parts)
    if not ocr_text.strip():
        logger.info("OCR yielded no text, trying PyPDF2 fallback extraction...")
        try:
            fallback_text = extract_text_from_pdf(file_bytes)
            if fallback_text.strip():
                logger.info("PyPDF2 fallback extracted text successfully.")
                return fallback_text
        except Exception as fb_err:
            logger.warning(f"PyPDF2 fallback also failed: {fb_err}")

    return ocr_text


def extract_text(file_bytes: bytes, file_type: str, is_scanned: bool = False) -> str:
    """Extract text from a file based on its type.

    Args:
        file_bytes: Raw file content
        file_type: File extension (e.g., '.pdf', '.docx')
        is_scanned: If True, use OCR for PDF extraction (for scanned documents)
    """
    if file_type == ".pdf":
        if is_scanned:
            return extract_text_from_scanned_pdf(file_bytes)
        return extract_text_from_pdf(file_bytes)
    elif file_type == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
